"""Generate all parser test fixtures. Idempotent."""
import json
import mailbox
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
import pandas as pd
from docx import Document

FIXTURE_DIR = Path(__file__).parent


def make_pdf() -> None:
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Inbound Lead SOP\nStep 1: collect contact info.\nStep 2: send document request.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Renewal SOP\nStep 1: pull declarations page.\nStep 2: email renewal quote.")
    doc.save(FIXTURE_DIR / "sop.pdf")
    doc.close()


def make_docx() -> None:
    doc = Document()
    doc.add_paragraph("Onboarding SOP")
    doc.add_paragraph("Step 1: verify lead identity.")
    doc.add_paragraph("Step 2: open file in Applied Epic.")
    doc.save(FIXTURE_DIR / "sop.docx")


def make_md() -> None:
    (FIXTURE_DIR / "notes.md").write_text(
        "# Producer Notes\n\nLeads waiting > 24h before first response.\nCSR manually copies CRM notes.\n"
    )


def make_txt() -> None:
    (FIXTURE_DIR / "notes.txt").write_text(
        "Discovery call summary.\nClient mentioned slow document collection.\nProducer follow-up inconsistent.\n"
    )


def make_vtt() -> None:
    (FIXTURE_DIR / "call.vtt").write_text(
        "WEBVTT\n\n"
        "00:00:01.000 --> 00:00:05.000\n"
        "Founder: our biggest issue is lead response time.\n\n"
        "00:00:06.000 --> 00:00:10.000\n"
        "CSR: we copy notes from email to HubSpot manually.\n"
    )


def make_srt() -> None:
    (FIXTURE_DIR / "call.srt").write_text(
        "1\n00:00:01,000 --> 00:00:05,000\nFounder: our biggest issue is lead response time.\n\n"
        "2\n00:00:06,000 --> 00:00:10,000\nCSR: we copy notes from email to HubSpot manually.\n"
    )


def make_csv() -> None:
    df = pd.DataFrame(
        [
            {"name": "Acme Corp", "email": "ops@acme.com", "stage": "awaiting_docs", "days_in_stage": 12},
            {"name": "Beta LLC", "email": "hi@beta.com", "stage": "new", "days_in_stage": 1},
            {"name": "Gamma Inc", "email": "info@gamma.com", "stage": "awaiting_docs", "days_in_stage": 30},
        ]
    )
    df.to_csv(FIXTURE_DIR / "leads.csv", index=False)


def make_xlsx() -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Leads"
    ws.append(["name", "email", "stage", "days_in_stage"])
    ws.append(["Acme Corp", "ops@acme.com", "awaiting_docs", 12])
    ws.append(["Beta LLC", "hi@beta.com", "new", 1])
    wb.save(FIXTURE_DIR / "leads.xlsx")


def make_mbox() -> None:
    path = FIXTURE_DIR / "inbox.mbox"
    if path.exists():
        path.unlink()
    mbox = mailbox.mbox(str(path))
    msg = mailbox.mboxMessage()
    msg["From"] = "lead@acme.com"
    msg["To"] = "csr@agency.com"
    msg["Subject"] = "Need a quote"
    msg["Message-ID"] = "<msg-001@acme.com>"
    msg.set_payload("Hi, we need a commercial liability quote ASAP. Please send the document list.")
    mbox.add(msg)
    mbox.close()


def make_json() -> None:
    payload = {
        "contacts": [
            {"id": "c1", "name": "Acme Corp", "last_touch_days": 12, "stage": "awaiting_docs"},
            {"id": "c2", "name": "Beta LLC", "last_touch_days": 1, "stage": "new"},
        ]
    }
    (FIXTURE_DIR / "crm.json").write_text(json.dumps(payload, indent=2))


def main() -> None:
    make_pdf()
    make_docx()
    make_md()
    make_txt()
    make_vtt()
    make_srt()
    make_csv()
    make_xlsx()
    make_mbox()
    make_json()
    print(f"Fixtures written to {FIXTURE_DIR}")


if __name__ == "__main__":
    main()
